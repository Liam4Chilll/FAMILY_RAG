"""Chargement et parsing des documents multi-formats."""

import os
import re
import email
import chardet
from pathlib import Path
from typing import List, Optional, Dict
from dataclasses import dataclass
from datetime import datetime

from langchain.schema import Document
from pypdf import PdfReader
from docx import Document as DocxDocument
import pytesseract
from PIL import Image

from config import get_settings


@dataclass
class LoadedDocument:
    """Document chargé avec métadonnées."""
    filename: str
    content: str
    file_type: str
    size_bytes: int
    error: Optional[str] = None


class DocumentLoader:
    """Chargeur de documents multi-formats."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.md', '.docx', '.eml', '.jpg', '.jpeg', '.png'}
    
    def __init__(self):
        self.settings = get_settings()
        self.data_dir = Path(self.settings.data_dir)
    
    def list_files(self) -> List[dict]:
        """Liste tous les fichiers supportés dans le dossier data."""
        files = []
        if not self.data_dir.exists():
            return files
        
        for file_path in self.data_dir.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                files.append({
                    'name': file_path.name,
                    'path': str(file_path.relative_to(self.data_dir)),
                    'type': file_path.suffix.lower()[1:],
                    'size': file_path.stat().st_size
                })
        
        return sorted(files, key=lambda x: x['name'].lower())

    def _extract_date(self, content: str, file_path: Path) -> Optional[str]:
        """Extrait une date du contenu ou du nom de fichier.

        Recherche des patterns de dates dans le contenu et le nom de fichier.
        Retourne la date au format ISO (YYYY-MM-DD) ou None.
        """
        # Pattern de dates courants
        date_patterns = [
            r'(\d{4}[-/]\d{2}[-/]\d{2})',  # YYYY-MM-DD ou YYYY/MM/DD
            r'(\d{2}[-/]\d{2}[-/]\d{4})',  # DD-MM-YYYY ou DD/MM/YYYY
            r'(\d{1,2}\s+(?:janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+\d{4})',  # FR
            r'(\d{1,2}\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+\d{4})',  # EN
        ]

        # Chercher dans le nom de fichier d'abord
        filename = file_path.stem
        for pattern in date_patterns[:2]:  # Uniquement formats numériques pour filename
            match = re.search(pattern, filename, re.IGNORECASE)
            if match:
                date_str = match.group(1)
                try:
                    # Normaliser au format ISO
                    if '/' in date_str or '-' in date_str:
                        parts = re.split(r'[-/]', date_str)
                        if len(parts[0]) == 4:  # YYYY-MM-DD
                            return f"{parts[0]}-{parts[1]}-{parts[2]}"
                        else:  # DD-MM-YYYY
                            return f"{parts[2]}-{parts[1]}-{parts[0]}"
                except:
                    pass

        # Chercher dans le contenu (premières 500 lignes)
        content_preview = content[:5000] if len(content) > 5000 else content
        for pattern in date_patterns:
            match = re.search(pattern, content_preview, re.IGNORECASE)
            if match:
                return match.group(1)

        # Dernière option : date de modification du fichier
        try:
            mtime = file_path.stat().st_mtime
            return datetime.fromtimestamp(mtime).strftime('%Y-%m-%d')
        except:
            return None

    def _extract_year(self, file_path: Path, content: str = "") -> Optional[int]:
        """Extrait l'année du nom de fichier ou du contenu."""
        # Chercher année dans le nom de fichier
        filename = file_path.stem
        year_match = re.search(r'(20\d{2}|19\d{2})', filename)
        if year_match:
            return int(year_match.group(1))

        # Chercher dans le contenu
        if content:
            content_preview = content[:1000]
            year_match = re.search(r'(20\d{2}|19\d{2})', content_preview)
            if year_match:
                return int(year_match.group(1))

        # Année de modification du fichier
        try:
            mtime = file_path.stat().st_mtime
            return datetime.fromtimestamp(mtime).year
        except:
            return None

    def _classify_document(self, content: str) -> str:
        """Classifie le type de document selon son contenu.

        Types possibles : contrat, recette, email, facture, note, autre
        """
        content_lower = content.lower()[:2000]  # Analyser le début

        # Mots-clés par catégorie
        keywords = {
            'contrat': ['contrat', 'partie', 'signataire', 'article', 'clause', 'conditions générales', 'prêt', 'emprunt'],
            'recette': ['ingrédients', 'recette', 'cuisson', 'préparation', 'four', 'ml', 'grammes', 'servir'],
            'facture': ['facture', 'montant', 'ttc', 'tva', 'total', 'paiement', 'numéro de facture'],
            'email': ['objet:', 'de:', 'à:', 'sujet:', 'from:', 'to:', 'subject:'],
            'formulaire': ['formulaire', 'nom:', 'prénom:', 'adresse:', 'date de naissance', 'signature'],
            'note': ['note', 'mémo', 'remarque', 'à faire', 'todo', 'rappel'],
        }

        # Compter les occurrences
        scores = {}
        for doc_type, words in keywords.items():
            score = sum(1 for word in words if word in content_lower)
            scores[doc_type] = score

        # Retourner le type avec le meilleur score
        if max(scores.values()) > 0:
            return max(scores, key=scores.get)
        return 'autre'

    def _extract_author(self, content: str) -> Optional[str]:
        """Extrait l'auteur/expéditeur du document."""
        content_preview = content[:1000]

        # Patterns pour extraire l'auteur
        patterns = [
            r'(?:De|From):\s*([^\n<]+)',  # Email
            r'(?:Auteur|Author):\s*([^\n]+)',  # Document
            r'(?:Par|By):\s*([^\n]+)',  # Note
            r'(?:Signé|Signed):\s*([^\n]+)',  # Contrat
        ]

        for pattern in patterns:
            match = re.search(pattern, content_preview, re.IGNORECASE)
            if match:
                author = match.group(1).strip()
                # Nettoyer (enlever email entre <>)
                author = re.sub(r'<[^>]+>', '', author).strip()
                if author:
                    return author[:100]  # Limiter la longueur

        return None

    def load_all(self) -> List[Document]:
        """Charge tous les documents du dossier data avec métadonnées enrichies."""
        documents = []

        for file_info in self.list_files():
            file_path = self.data_dir / file_info['path']
            loaded = self._load_file(file_path)

            if loaded.content and not loaded.error:
                # Extraire métadonnées enrichies
                doc_date = self._extract_date(loaded.content, file_path)
                doc_year = self._extract_year(file_path, loaded.content)
                doc_type = self._classify_document(loaded.content)
                doc_author = self._extract_author(loaded.content)

                # Créer le document avec métadonnées complètes
                metadata = {
                    'source': loaded.filename,
                    'file_type': loaded.file_type,
                    'size_bytes': loaded.size_bytes,
                    'date': doc_date,
                    'year': doc_year,
                    'doc_type': doc_type,
                    'author': doc_author
                }

                # Supprimer les None pour alléger
                metadata = {k: v for k, v in metadata.items() if v is not None}

                doc = Document(
                    page_content=loaded.content,
                    metadata=metadata
                )
                documents.append(doc)

                # Log des métadonnées extraites
                print(f"[Loader] {loaded.filename} → type:{doc_type}, année:{doc_year}, auteur:{doc_author or 'N/A'}")

        return documents

    def load_specific(self, file_paths: List[str]) -> List[Document]:
        """Charge uniquement les fichiers spécifiés.

        Args:
            file_paths: Liste des chemins relatifs des fichiers à charger

        Returns:
            Liste de Documents Langchain
        """
        documents = []

        for relative_path in file_paths:
            file_path = self.data_dir / relative_path

            if not file_path.exists():
                print(f"[Loader] ⚠️ Fichier introuvable : {relative_path}")
                continue

            if not file_path.is_file():
                continue

            if file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
                print(f"[Loader] ⚠️ Extension non supportée : {relative_path}")
                continue

            # Charger le fichier
            loaded = self._load_file(file_path)

            if loaded.error:
                print(f"[Loader] ❌ Erreur chargement {loaded.filename}: {loaded.error}")
            elif loaded.content:
                # Extraire métadonnées enrichies
                doc_date = self._extract_date(loaded.content, file_path)
                doc_year = self._extract_year(file_path, loaded.content)
                doc_type = self._classify_document(loaded.content)
                doc_author = self._extract_author(loaded.content)

                # Créer le document avec métadonnées complètes
                metadata = {
                    'source': loaded.filename,
                    'file_type': loaded.file_type,
                    'size_bytes': loaded.size_bytes,
                    'date': doc_date,
                    'year': doc_year,
                    'doc_type': doc_type,
                    'author': doc_author
                }

                # Supprimer les None pour alléger
                metadata = {k: v for k, v in metadata.items() if v is not None}

                doc = Document(
                    page_content=loaded.content,
                    metadata=metadata
                )
                documents.append(doc)

                # Log des métadonnées extraites
                print(f"[Loader] {loaded.filename} → type:{doc_type}, année:{doc_year}, auteur:{doc_author or 'N/A'}")

        print(f"[Loader] {len(documents)}/{len(file_paths)} fichiers chargés avec succès")
        return documents
    
    def _load_file(self, file_path: Path) -> LoadedDocument:
        """Charge un fichier selon son extension."""
        ext = file_path.suffix.lower()
        size = file_path.stat().st_size
        
        try:
            if ext == '.pdf':
                content = self._load_pdf(file_path)
            elif ext in {'.txt', '.md'}:
                content = self._load_text(file_path)
            elif ext == '.docx':
                content = self._load_docx(file_path)
            elif ext == '.eml':
                content = self._load_eml(file_path)
            elif ext in {'.jpg', '.jpeg', '.png'}:
                content = self._load_image(file_path)
            else:
                return LoadedDocument(
                    filename=file_path.name,
                    content='',
                    file_type=ext[1:],
                    size_bytes=size,
                    error=f'Extension non supportée: {ext}'
                )
            
            return LoadedDocument(
                filename=file_path.name,
                content=content,
                file_type=ext[1:],
                size_bytes=size
            )
            
        except Exception as e:
            return LoadedDocument(
                filename=file_path.name,
                content='',
                file_type=ext[1:],
                size_bytes=size,
                error=str(e)
            )
    
    def _load_pdf(self, file_path: Path) -> str:
        """Extrait le texte d'un PDF."""
        reader = PdfReader(file_path)
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return '\n\n'.join(texts)
    
    def _load_text(self, file_path: Path) -> str:
        """Charge un fichier texte avec détection d'encodage."""
        raw = file_path.read_bytes()
        detected = chardet.detect(raw)
        encoding = detected.get('encoding', 'utf-8') or 'utf-8'
        return raw.decode(encoding, errors='replace')
    
    def _load_docx(self, file_path: Path) -> str:
        """Extrait le texte d'un fichier Word."""
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n\n'.join(paragraphs)
    
    def _load_eml(self, file_path: Path) -> str:
        """Extrait le contenu d'un email."""
        raw = file_path.read_bytes()
        msg = email.message_from_bytes(raw)
        
        parts = []
        
        # En-têtes
        if msg['subject']:
            parts.append(f"Sujet: {msg['subject']}")
        if msg['from']:
            parts.append(f"De: {msg['from']}")
        if msg['to']:
            parts.append(f"À: {msg['to']}")
        if msg['date']:
            parts.append(f"Date: {msg['date']}")
        
        parts.append('')  # Ligne vide
        
        # Corps
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == 'text/plain':
                    payload = part.get_payload(decode=True)
                    if payload:
                        charset = part.get_content_charset() or 'utf-8'
                        parts.append(payload.decode(charset, errors='replace'))
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                charset = msg.get_content_charset() or 'utf-8'
                parts.append(payload.decode(charset, errors='replace'))
        
        return '\n'.join(parts)
    
    def _load_image(self, file_path: Path) -> str:
        """Extrait le texte d'une image via OCR (Tesseract)."""
        image = Image.open(file_path)
        # Utilise français + anglais pour la reconnaissance
        text = pytesseract.image_to_string(image, lang='fra+eng')
        return text.strip()
